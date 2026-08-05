"""
Document & Content Extractors for Crawl-ETL
"""

import io
import re
import json
import logging
from typing import Dict, Any, List, Tuple
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse

logger = logging.getLogger("CrawlETL.Extractors")

class HTMLExtractor:
    @staticmethod
    def extract(html_content: str, base_url: str) -> Dict[str, Any]:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script, style, nav, footer noise for main text extraction
        for element in soup(["script", "style", "nav", "footer", "iframe", "noscript"]):
            element.decompose()

        # Title
        title_tag = soup.find('title')
        title = title_tag.get_text().strip() if title_tag else base_url

        # Clean text
        text = soup.get_text(separator='\n')
        cleaned_lines = [line.strip() for line in text.splitlines() if line.strip()]
        cleaned_text = '\n'.join(cleaned_lines)

        # Extract links
        links = []
        for a in soup.find_all('a', href=True):
            href = a['href'].strip()
            if href and not href.startswith(('javascript:', 'mailto:', 'tel:', '#')):
                full_url = urljoin(base_url, href)
                links.append(full_url)
        links = list(set(links))

        # Extract media URLs
        media_urls = []
        # Images
        for img in soup.find_all('img', src=True):
            src = urljoin(base_url, img['src'])
            media_urls.append({"type": "image", "url": src, "alt": img.get('alt', '')})

        # Videos
        for video in soup.find_all(['video', 'source']):
            src = video.get('src')
            if src:
                media_urls.append({"type": "video", "url": urljoin(base_url, src)})

        # Audio / Podcast
        for audio in soup.find_all(['audio', 'source']):
            src = audio.get('src')
            if src:
                media_urls.append({"type": "audio", "url": urljoin(base_url, src)})

        # Tables
        tables = []
        for table in soup.find_all('table'):
            table_data = []
            for row in table.find_all('tr'):
                cols = [col.get_text().strip() for col in row.find_all(['td', 'th'])]
                if cols:
                    table_data.append(cols)
            if table_data:
                tables.append(table_data)

        # OpenGraph / JSON-LD Metadata
        metadata = {}
        for meta in soup.find_all('meta'):
            if meta.get('property') and meta.get('content'):
                metadata[meta['property']] = meta['content']
            elif meta.get('name') and meta.get('content'):
                metadata[meta['name']] = meta['content']

        return {
            "title": title,
            "cleaned_text": cleaned_text,
            "links": links,
            "media_urls": media_urls,
            "tables": tables,
            "metadata": metadata
        }

class DocumentExtractor:
    @staticmethod
    def extract_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text and tables from PDF files"""
        extracted_text = []
        tables = []
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    txt = page.extract_text()
                    if txt:
                        extracted_text.append(f"--- Page {page_idx + 1} ---\n" + txt)
                    page_tables = page.extract_tables()
                    for t in page_tables:
                        cleaned_t = [[cell.strip() if cell else "" for cell in row] for row in t]
                        tables.append(cleaned_t)
        except Exception as e:
            logger.warning(f"Fallback pdf extraction: {e}")
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page_idx, page in enumerate(reader.pages):
                    txt = page.extract_text()
                    if txt:
                        extracted_text.append(f"--- Page {page_idx + 1} ---\n" + txt)
            except Exception as ex:
                logger.error(f"Failed PDF extraction: {ex}")

        full_text = "\n\n".join(extracted_text)
        return {
            "title": "PDF Document",
            "cleaned_text": full_text,
            "tables": tables,
            "metadata": {"type": "pdf", "page_count": len(extracted_text)}
        }

    @staticmethod
    def extract_docx(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text and tables from DOCX files"""
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            
            tables = []
            for table in doc.tables:
                t_data = []
                for row in table.rows:
                    cols = [cell.text.strip() for cell in row.cells]
                    t_data.append(cols)
                tables.append(t_data)
                
            return {
                "title": "Word Document",
                "cleaned_text": full_text,
                "tables": tables,
                "metadata": {"type": "docx"}
            }
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return {"title": "Word Document", "cleaned_text": "", "tables": [], "metadata": {}}

    @staticmethod
    def extract_excel(file_bytes: bytes) -> Dict[str, Any]:
        """Extract spreadsheets from XLSX/CSV files"""
        try:
            import pandas as pd
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
            all_text = []
            tables = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                df_filled = df.fillna("")
                all_text.append(f"=== Sheet: {sheet_name} ===\n" + df_filled.to_string())
                tables.append([df.columns.tolist()] + df_filled.values.tolist())
            return {
                "title": "Excel Spreadsheet",
                "cleaned_text": "\n\n".join(all_text),
                "tables": tables,
                "metadata": {"sheets": excel_file.sheet_names}
            }
        except Exception as e:
            logger.error(f"Excel extraction error: {e}")
            return {"title": "Excel Spreadsheet", "cleaned_text": "", "tables": [], "metadata": {}}
